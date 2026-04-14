import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
import re
from io import BytesIO

# ================================================================
#  WEIGHTS
# ================================================================
WEIGHTS = {"skills": 0.40, "experience": 0.40, "education": 0.20}

# ================================================================
#  EDUCATION HIERARCHY
# ================================================================
EDUCATION_RANK = {
    "phd": 5, "ph.d": 5, "doctorate": 5, "doctoral": 5, "d.phil": 5,
    "master": 4, "mba": 4, "m.sc": 4, "msc": 4, "m.tech": 4, "mtech": 4,
    "m.e.": 4, "m.s.": 4, "pg diploma": 4, "post graduate diploma": 4,
    "postgraduate diploma": 4, "post-graduate": 4,
    "bachelor": 3, "b.sc": 3, "bsc": 3, "b.tech": 3, "btech": 3,
    "b.e.": 3, "b.a.": 3, "b.com": 3, "bcom": 3, "bca": 3, "b.ca": 3,
    "undergraduate": 3, "graduation": 3, "graduate": 3,
    "diploma": 2, "polytechnic": 2,
    "12th": 1, "hsc": 1, "intermediate": 1, "higher secondary": 1,
    "10th": 0, "ssc": 0, "matriculation": 0,
}

EDUCATION_LABEL = {
    5: "PhD / Doctorate",
    4: "Master's / PG Diploma",
    3: "Bachelor's Degree",
    2: "Diploma",
    1: "12th / HSC",
    0: "10th / SSC",
}

# ================================================================
#  CURATED SKILL PHRASES
#  Only these will be extracted as skills — prevents noise words
#  from JD being treated as skills (e.g. "account", "abilities")
# ================================================================
KNOWN_SKILLS = [
    # Programming languages
    "python", "java", "javascript", "typescript", "c++", "c#", "ruby",
    "go", "rust", "swift", "kotlin", "php", "matlab", "scala", "perl",
    "r programming", "vba", "shell scripting", "bash",
    # Web / Frontend
    "html", "css", "react", "angular", "vue", "node.js", "nodejs",
    "jquery", "bootstrap", "tailwind",
    # Backend / Frameworks
    "django", "flask", "fastapi", "spring boot", "express", "laravel",
    "asp.net", "ruby on rails",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "oracle", "sqlite",
    "redis", "cassandra", "dynamodb", "ms sql", "sql server",
    # Cloud / DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
    "ansible", "jenkins", "ci/cd", "devops", "linux", "unix",
    "git", "github", "gitlab", "bitbucket",
    # Data / ML / AI
    "machine learning", "deep learning", "nlp", "data science",
    "data analysis", "data analytics", "data engineering",
    "tensorflow", "pytorch", "scikit-learn", "keras", "opencv",
    "pandas", "numpy", "scipy", "matplotlib", "seaborn",
    "spark", "hadoop", "kafka", "airflow", "dbt",
    # BI / Reporting
    "power bi", "tableau", "looker", "qlik", "excel", "pivot tables",
    "google analytics", "data studio",
    # Finance / Accounting specific
    "tally", "sap", "sap fico", "sap mm", "sap sd", "erp",
    "quickbooks", "zoho books", "busy accounting", "marg erp",
    "gst", "tds", "income tax", "taxation", "audit", "auditing",
    "financial reporting", "financial analysis", "financial modelling",
    "accounts payable", "accounts receivable", "bookkeeping",
    "balance sheet", "profit and loss", "cash flow", "budgeting",
    "forecasting", "cost accounting", "statutory compliance",
    "payroll", "bank reconciliation", "ledger", "journal entries",
    "ifrs", "gaap", "ind as",
    # General Tech
    "rest api", "graphql", "microservices", "agile", "scrum",
    "jira", "confluence", "figma", "photoshop", "ms office",
    "word", "powerpoint", "outlook", "sharepoint",
    # Soft skills
    "project management", "team management", "client management",
    "stakeholder management", "vendor management",
    "communication", "leadership", "negotiation",
    "problem solving", "critical thinking", "cross-functional",
    "presentation skills", "interpersonal skills",
    # Marketing / Sales
    "digital marketing", "seo", "sem", "social media marketing",
    "content marketing", "email marketing", "crm",
    "salesforce", "hubspot", "google ads", "facebook ads",
    # Engineering / Manufacturing
    "autocad", "solidworks", "catia", "ansys",
    "quality control", "quality assurance", "six sigma", "lean",
    "supply chain", "logistics", "procurement", "inventory management",
]

KNOWN_SKILLS_SORTED = sorted(set(KNOWN_SKILLS), key=len, reverse=True)

# ================================================================
#  WORK SECTION MARKERS
# ================================================================
WORK_SECTION_HEADERS = [
    "work experience", "professional experience", "employment history",
    "career history", "work history", "positions held",
    "professional background", "experience",
]

SECTION_END_MARKERS = [
    "education", "qualification", "academic", "certification", "skills",
    "projects", "awards", "hobbies", "references", "declaration",
    "achievements", "extra", "volunteer", "publication", "training",
    "courses", "interests", "languages",
]


# ================================================================
#  TEXT EXTRACTION
# ================================================================

def extract_text(file) -> str:
    if file.name.lower().endswith(".pdf"):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text
    else:
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)


# ================================================================
#  NAME EXTRACTION
# ================================================================

def extract_name(text: str, file_name: str) -> str:
    lines = [l.strip() for l in text.split("\n") if l.strip()][:10]
    for line in lines:
        words = line.split()
        if (2 <= len(words) <= 5
                and all(re.match(r"^[A-Za-z.'\-]+$", w) for w in words)
                and not any(kw in line.lower() for kw in [
                    "resume", "curriculum", "vitae", "profile", "objective",
                    "summary", "email", "phone", "address", "linkedin",
                    "mobile", "contact", "experienced", "professional",
                ])):
            return line.title()
    name = re.sub(r"\.(pdf|docx)$", "", file_name, flags=re.IGNORECASE)
    name = re.sub(r"[_\-]", " ", name)
    name = re.sub(r"(resume|cv|curriculum|vitae)", "", name, flags=re.IGNORECASE)
    name = re.sub(r"\d+", "", name)
    return name.strip().title() or "Unknown"


# ================================================================
#  JD PARSING
# ================================================================

def extract_jd_skills(jd_text: str) -> list:
    """Only match from KNOWN_SKILLS — prevents noise words being treated as skills."""
    jd_lower = jd_text.lower()
    found = []
    seen = set()
    for phrase in KNOWN_SKILLS_SORTED:
        if len(phrase) <= 4:
            pattern = r'\b' + re.escape(phrase) + r'\b'
            if re.search(pattern, jd_lower) and phrase not in seen:
                found.append(phrase)
                seen.add(phrase)
        else:
            if phrase in jd_lower and phrase not in seen:
                found.append(phrase)
                seen.add(phrase)
    return found


def extract_required_experience(jd_text: str) -> float:
    jd_lower = jd_text.lower()
    patterns = [
        r"(\d+)\s*\+\s*years?\s+of\s+(?:relevant\s+)?(?:work\s+)?experience",
        r"minimum\s+(?:of\s+)?(\d+)\s+years?",
        r"at\s+least\s+(\d+)\s+years?",
        r"(\d+)\s+to\s+\d+\s+years?\s+of\s+experience",
        r"(\d+)\s*[-–]\s*\d+\s+years?\s+of\s+experience",
        r"(\d+)\s*\+?\s*years?\s+of\s+(?:relevant\s+)?(?:work\s+)?experience",
        r"experience\s*[:\-]?\s*(\d+)\s*\+?\s*years?",
        r"(\d+)\s+years?\s+experience",
    ]
    for pat in patterns:
        m = re.search(pat, jd_lower)
        if m:
            val = float(m.group(1))
            if 0 < val <= 40:
                return val
    return 3.0


def extract_required_education(jd_text: str) -> int:
    jd_lower = jd_text.lower()
    best = 0
    for keyword, rank in sorted(EDUCATION_RANK.items(), key=lambda x: -x[1]):
        if keyword in jd_lower and rank > best:
            best = rank
    return best if best > 0 else 3


# ================================================================
#  RESUME — SKILLS
# ================================================================

def extract_candidate_skills(resume_text: str) -> list:
    resume_lower = resume_text.lower()
    found = []
    seen = set()
    for phrase in KNOWN_SKILLS_SORTED:
        if len(phrase) <= 4:
            pattern = r'\b' + re.escape(phrase) + r'\b'
            if re.search(pattern, resume_lower) and phrase not in seen:
                found.append(phrase)
                seen.add(phrase)
        else:
            if phrase in resume_lower and phrase not in seen:
                found.append(phrase)
                seen.add(phrase)
    return found


# ================================================================
#  RESUME — EXPERIENCE (FIXED)
# ================================================================

def isolate_work_section(text: str) -> str:
    """
    Isolate just the work experience section to avoid summing education years.
    Falls back to full text if section headers not found.
    """
    lines = text.split("\n")
    start_idx = None
    end_idx = len(lines)

    for i, line in enumerate(lines):
        line_lower = line.strip().lower()
        # Must be a short header line (not a content line with dates)
        if (any(line_lower == h or line_lower.startswith(h)
                for h in WORK_SECTION_HEADERS)
                and not re.search(r'\d{4}', line)
                and len(line.strip()) < 40):
            start_idx = i
            break

    if start_idx is None:
        return text  # fallback

    for i in range(start_idx + 1, len(lines)):
        line_lower = lines[i].strip().lower()
        if (any(line_lower == m or line_lower.startswith(m)
                for m in SECTION_END_MARKERS)
                and not re.search(r'\d{4}', lines[i])
                and len(lines[i].strip()) < 40):
            end_idx = i
            break

    return "\n".join(lines[start_idx:end_idx])


def extract_experience_years(resume_text: str) -> float:
    """
    Three-strategy experience extractor.
    Strategy 1: Explicit total experience statement (most reliable).
    Strategy 2: Sum of date ranges in work section ONLY.
    Strategy 3: Year span in work section ONLY.
    """
    current_year = 2025

    # Strategy 1: Explicit statement
    explicit_patterns = [
        (r'(\d+)\s+years?\s+(\d+)\s+months?\s+(?:of\s+)?(?:total\s+)?experience', True),
        (r'(\d+\.?\d*)\s*\+?\s*years?\s+of\s+(?:total\s+)?(?:work\s+)?(?:professional\s+)?experience', False),
        (r'total\s+(?:work\s+)?experience\s*[:\-]?\s*(\d+\.?\d*)\s*years?', False),
        (r'(\d+\.?\d*)\s*years?\s+of\s+(?:total\s+)?(?:work|professional|industry)\s+experience', False),
        (r'experience\s*[:\-]\s*(\d+\.?\d*)\s*years?', False),
    ]
    for pat, two_groups in explicit_patterns:
        m = re.search(pat, resume_text, re.IGNORECASE)
        if m:
            val = float(m.group(1)) + float(m.group(2)) / 12 if two_groups else float(m.group(1))
            if 0 < val <= 45:
                return round(val, 1)

    # Strategy 2: Date ranges in WORK SECTION only
    work_text = isolate_work_section(resume_text)

    range_pattern = r'''
        (?:
            (?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|
               Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)
            \.?\s+
        )?
        ((?:19|20)\d{2})
        \s*[-–—to/]+\s*
        (?:
            (?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|
               Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)
            \.?\s+
        )?
        ((?:19|20)\d{2}|present|current|now|till\s*date|to\s*date|date)
    '''
    matches = re.findall(range_pattern, work_text, re.IGNORECASE | re.VERBOSE)

    if matches:
        seen_ranges = set()
        total_months = 0.0
        for start_s, end_s in matches:
            start = int(start_s)
            end = current_year if re.match(r'[a-z]', end_s.strip(), re.IGNORECASE) else int(end_s)
            if not (1970 <= start <= current_year and start <= end <= current_year + 1):
                continue
            key = (start, end)
            if key in seen_ranges:
                continue
            seen_ranges.add(key)
            # Cap each single job at 15 years to prevent education years bleeding in
            total_months += min((end - start) * 12, 180)

        if total_months > 0:
            return min(round(total_months / 12, 1), 45.0)

    # Strategy 3: Year span in work section only
    years_in_work = [
        int(y) for y in re.findall(r'\b((?:19|20)\d{2})\b', work_text)
        if 1970 <= int(y) <= current_year
    ]
    if len(years_in_work) >= 2:
        span = max(years_in_work) - min(years_in_work)
        if span <= 40:
            return float(span)

    return 0.0


# ================================================================
#  RESUME — EDUCATION
# ================================================================

def extract_education(resume_text: str) -> tuple:
    text = resume_text.lower()
    best_rank = -1
    for keyword, rank in sorted(EDUCATION_RANK.items(), key=lambda x: -x[1]):
        if keyword in text and rank > best_rank:
            best_rank = rank
    if best_rank == -1:
        return (0, "Not Found")
    return (best_rank, EDUCATION_LABEL.get(best_rank, "Other"))


# ================================================================
#  SCORING
# ================================================================

def score_skills(candidate_skills: list, jd_skills: list) -> tuple:
    if not jd_skills:
        return 1.0, [], []
    jd_set = set(jd_skills)
    candidate_set = set(candidate_skills)
    matched = sorted(jd_set & candidate_set)
    missing = sorted(jd_set - candidate_set)
    return len(matched) / len(jd_set), matched, missing


def score_experience(candidate_years: float, required_years: float) -> float:
    if required_years <= 0:
        return 1.0
    return min(candidate_years / required_years, 1.0)


def score_education(candidate_rank: int, required_rank: int) -> float:
    if required_rank <= 0:
        return 1.0
    if candidate_rank >= required_rank:
        return 1.0
    diff = required_rank - candidate_rank
    if diff == 1:
        return 0.6
    elif diff == 2:
        return 0.3
    return 0.0


def build_strengths_gaps(
    matched_skills, missing_skills,
    candidate_years, required_years,
    candidate_edu_rank, required_edu_rank, candidate_edu_label
) -> tuple:
    strengths, gaps = [], []
    seen_s, seen_g = set(), set()

    def add_strength(s):
        k = s.lower()
        if k not in seen_s and len(strengths) < 3:
            seen_s.add(k)
            strengths.append(s)

    def add_gap(g):
        k = g.lower()
        if k not in seen_g and len(gaps) < 3:
            seen_g.add(k)
            gaps.append(g)

    # Skills
    if matched_skills:
        add_strength(f"Matches key skills: {', '.join(matched_skills[:5])}")
    if missing_skills:
        clean_missing = [s for s in missing_skills if s not in matched_skills][:4]
        if clean_missing:
            add_gap(f"Missing skills: {', '.join(clean_missing)}")

    # Experience
    if required_years > 0:
        if candidate_years == 0:
            add_gap("Experience could not be determined from resume")
        elif candidate_years >= required_years * 1.25:
            add_strength(
                f"Exceeds experience requirement "
                f"({candidate_years:.1f} yrs vs {required_years:.0f} required)"
            )
        elif candidate_years >= required_years:
            add_strength(f"Meets experience requirement ({candidate_years:.1f} yrs)")
        else:
            add_gap(
                f"Below required experience "
                f"({candidate_years:.1f} yrs vs {required_years:.0f} needed)"
            )

    # Education
    req_label = EDUCATION_LABEL.get(required_edu_rank, "required level")
    if candidate_edu_rank > required_edu_rank:
        add_strength(f"Education exceeds requirement ({candidate_edu_label})")
    elif candidate_edu_rank == required_edu_rank:
        add_strength(f"Meets education requirement ({candidate_edu_label})")
    else:
        add_gap(
            f"Education below requirement "
            f"(has {candidate_edu_label}, needs {req_label})"
        )

    return strengths, gaps


# ================================================================
#  FULL ANALYSIS
# ================================================================

def analyze_candidate(resume_text, file_name, jd_skills, required_years, required_edu_rank):
    name = extract_name(resume_text, file_name)
    candidate_skills = extract_candidate_skills(resume_text)
    candidate_years = extract_experience_years(resume_text)
    candidate_edu_rank, candidate_edu_label = extract_education(resume_text)

    skill_score, matched, missing = score_skills(candidate_skills, jd_skills)
    exp_score = score_experience(candidate_years, required_years)
    edu_score = score_education(candidate_edu_rank, required_edu_rank)

    final_score = round(
        (skill_score * WEIGHTS["skills"]
         + exp_score * WEIGHTS["experience"]
         + edu_score * WEIGHTS["education"]) * 100
    )

    strengths, gaps = build_strengths_gaps(
        matched, missing,
        candidate_years, required_years,
        candidate_edu_rank, required_edu_rank, candidate_edu_label,
    )

    return {
        "file_name": file_name,
        "name": name,
        "score": final_score,
        "experience_years": candidate_years,
        "education_label": candidate_edu_label,
        "education_rank": candidate_edu_rank,
        "strengths": strengths,
        "gaps": gaps,
        "matched_skills": matched,
        "missing_skills": missing,
    }


# ================================================================
#  REPORT GENERATION
# ================================================================

def generate_report(top_candidates, required_years, required_edu_rank):
    doc = Document()
    title = doc.add_heading("Top Candidates Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    req_edu_label = EDUCATION_LABEL.get(required_edu_rank, "Not specified")
    meta = doc.add_paragraph(
        f"Weights — Skills 40% · Experience 40% · Education 20%   |   "
        f"Required experience: {required_years:.0f} yrs   |   "
        f"Required education: {req_edu_label}"
    )
    meta.runs[0].italic = True
    doc.add_paragraph("")

    for i, c in enumerate(top_candidates, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {c['name']}  —  Match Score: {c['score']}%")
        run.bold = True
        run.font.size = Pt(12)

        exp_display = (
            f"{c['experience_years']:.1f} yrs"
            if c['experience_years'] > 0 else "Not found"
        )
        doc.add_paragraph(
            f"File: {c['file_name']}   |   "
            f"Experience: {exp_display}   |   "
            f"Highest Education: {c['education_label']}"
        )

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "✔ Strengths"
        hdr[1].text = "✘ Gaps"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        strengths = c.get("strengths", [])
        gaps = c.get("gaps", [])
        max_len = max(len(strengths), len(gaps), 1)
        for j in range(max_len):
            row = table.add_row().cells
            row[0].text = strengths[j] if j < len(strengths) else ""
            row[1].text = gaps[j] if j < len(gaps) else ""

        doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ================================================================
#  STREAMLIT UI
# ================================================================

st.set_page_config(page_title="Rule-Based Hiring Assistant", page_icon="🧑‍💼")
st.title("🧑‍💼 Rule-Based Hiring Assistant")
st.caption(
    "Deterministic scoring — same inputs always produce the same score.  "
    "Weights: Skills 40% · Experience 40% · Education 20%"
)

jd_file = st.file_uploader("📄 Upload Job Description (PDF or DOCX)", type=["pdf", "docx"])
jd_text = ""
if jd_file:
    jd_text = extract_text(jd_file)
    st.success("✅ Job Description uploaded")

    with st.expander("🔍 Verify what was parsed from your JD before analyzing"):
        jd_skills = extract_jd_skills(jd_text)
        req_years = extract_required_experience(jd_text)
        req_edu = extract_required_education(jd_text)

        col1, col2 = st.columns(2)
        with col1:
            st.metric("Required Experience", f"{req_years:.0f} years")
        with col2:
            st.metric("Required Education", EDUCATION_LABEL.get(req_edu, "N/A"))

        if jd_skills:
            st.write(f"**Skills extracted from JD ({len(jd_skills)}):**")
            st.write(", ".join(jd_skills))
        else:
            st.warning(
                "⚠️ No skills matched from the JD. "
                "The domain may not be in the skill list — all candidates will score 0 on skills."
            )

resume_files = st.file_uploader(
    "📂 Upload Resumes (PDF or DOCX)",
    type=["pdf", "docx"],
    accept_multiple_files=True,
)
if resume_files:
    st.success(f"✅ {len(resume_files)} resume(s) uploaded")

top_n = st.slider("Number of top candidates in report", 1, 20, 3)

if st.button("🔍 Analyze Candidates"):
    if not jd_text or not resume_files:
        st.warning("⚠️ Please upload both a Job Description and at least one Resume.")
    else:
        jd_skills = extract_jd_skills(jd_text)
        req_years = extract_required_experience(jd_text)
        req_edu = extract_required_education(jd_text)

        results = []
        progress = st.progress(0)
        for idx, file in enumerate(resume_files):
            resume_text = extract_text(file)
            result = analyze_candidate(
                resume_text, file.name, jd_skills, req_years, req_edu
            )
            results.append(result)
            progress.progress((idx + 1) / len(resume_files))

        sorted_results = sorted(results, key=lambda x: x["score"], reverse=True)
        top_candidates = sorted_results[:top_n]

        st.success("✅ Analysis complete")
        st.subheader("📊 All Candidates Ranked")

        for i, c in enumerate(sorted_results, 1):
            tag = "🏆" if i <= top_n else "  "
            exp_str = (
                f"{c['experience_years']:.1f} yrs"
                if c['experience_years'] > 0 else "N/A"
            )
            st.write(
                f"{tag} **#{i} {c['name']}** — Score: `{c['score']}%` | "
                f"Exp: `{exp_str}` | Edu: `{c['education_label']}`"
            )
            with st.expander(f"Details — {c['name']}"):
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**✔ Strengths**")
                    for s in (c["strengths"] or ["—"]):
                        st.write(f"• {s}")
                with col2:
                    st.write("**✘ Gaps**")
                    for g in (c["gaps"] or ["—"]):
                        st.write(f"• {g}")
                if c["matched_skills"]:
                    st.write(f"*Matched skills: {', '.join(c['matched_skills'])}*")
                if c["missing_skills"]:
                    st.write(f"*Missing JD skills: {', '.join(c['missing_skills'][:6])}*")

        report = generate_report(top_candidates, req_years, req_edu)
        st.download_button(
            label="📄 Download Top Candidates Report (.docx)",
            data=report,
            file_name="Top_Candidates_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
